from docx import Document
import xlsx_tools
from docx.shared import Inches

KEY_WIDTH_INCHES = Inches(1.0)
VALUE_WIDTH_INCHES = Inches(0.7)

test_typed_sheets = {
    '2g': ('Voice', 'Radio statistics'),
    '3g-voice': ('Voice', 'Radio statistics'),
    '3g-data': ('Data (1 of 2)', 'Data (2 of 2)'),
    '4g-data': ('Data (1 of 2)', 'Data (2 of 2)'),
    'volte': ('Volte',)
}

def compose_docx(target_name, excels_to_parse, tech_plus_test):
    document = Document()

    # PARSING FILE BUNDLE
    for date in sorted(excels_to_parse.keys()):
        document.add_paragraph(date).bold = True
        # print(test_typed_sheets[tech_plus_test])
        for sheet in test_typed_sheets[tech_plus_test]:
            parsed = xlsx_tools.parse_bundle(excels_to_parse[date], sheet)
            create_table(document, parsed)

    document.add_page_break()

    document.save(target_name)

def create_table(document, parsed):
    table = document.add_table(rows = 0, cols = 2, style = 'Light Grid Accent 6')
    for section in parsed[0].keys():
        # print(parsed[0].keys())
        if section == '__TYPE__':
            continue
        row_cells = table.add_row().cells
        row_cells[0].add_paragraph(text = section).bold = True
        #for operator
        row_cells = table.add_row().cells
        print(parsed[0]['__TYPE__'], len(parsed[0]['__TYPE__']))
        row_cells[1].add_paragraph(text = parsed[0]['__TYPE__']).bold = True
        for entry in parsed[0][section].keys():
            entries_row_cells = table.add_row().cells
            entries_row_cells[0].text = str(entry)
            entries_row_cells[1].text = str(round(parsed[0][section][entry], 2))

    for operator in parsed[1:]:
        col_cells = table.add_column(width = VALUE_WIDTH_INCHES).cells
        cell_counter = 0
        for section in operator.keys():
            if section == '__TYPE__':
                continue
            cell_counter = cell_counter + 1
            col_cells[cell_counter].text = str(operator['__TYPE__'])
            cell_counter = cell_counter + 1
            for entry in operator[section].keys():
                col_cells[cell_counter].text = str(round(operator[section][entry], 2))
                cell_counter = cell_counter + 1

    table.columns[0].width = KEY_WIDTH_INCHES
    table.columns[1].width = VALUE_WIDTH_INCHES

